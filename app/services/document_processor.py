from fastapi import UploadFile, HTTPException
from PyPDF2 import PdfReader
from docx import Document
import io
import magic

from app.logger import get_logger

# Get configured logger
logger = get_logger(__name__)

class DocumentProcessor:
    """Handles extraction of text from PDF and DOCX files"""
    
    # Maximum file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Supported MIME types
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
    }

    def extract_text(self, content: bytes, file_ext: str) -> str:
        """
        Extract text content from file bytes
        
        Args:
            content (bytes): File content as bytes
            file_ext (str): File extension including dot (e.g. '.pdf', '.docx')
        
        Returns:
            str: Extracted text content
        
        Raises:
            ValueError: If file format is not supported or validation fails
        """
        try:
            # Validate file size
            if len(content) > self.MAX_FILE_SIZE:
                logger.error(f"File too large ({len(content)/1024/1024:.2f} MB)")
                raise ValueError(f"File size exceeds maximum limit of {self.MAX_FILE_SIZE/1024/1024:.0f}MB")
                
            if len(content) == 0:
                logger.error("Empty file content received")
                raise ValueError("File content is empty")

            # Map file extensions to processing functions
            extension_map = {
                '.pdf': self._extract_from_pdf,
                '.docx': self._extract_from_docx
            }

            if file_ext.lower() not in extension_map:
                raise ValueError(f"Unsupported file extension: {file_ext}. Supported: {', '.join(extension_map.keys())}")

            # Get file size for logging
            file_size_kb = len(content) / 1024
            logger.info(f"Processing file content ({file_size_kb:.2f} KB)")
            
            return extension_map[file_ext.lower()](content, "document")
                
        except ValueError as e:
            raise e
        except Exception as e:
            logger.error(f"Error processing file content: {str(e)}")
            raise ValueError(f"Error processing file: {str(e)}")
    
    async def extract_text_from_upload(self, file: UploadFile) -> str:
        """
        Extract text content from uploaded PDF or DOCX file
        
        Args:
            file (UploadFile): Uploaded file object
        
        Returns:
            str: Extracted text content
        
        Raises:
            ValueError: If file format is not supported or validation fails
            HTTPException: If file processing fails
        """
        try:
            content = await file.read()
            
            # Validate file type using magic numbers
            mime_type = magic.from_buffer(content, mime=True)
            if mime_type not in self.SUPPORTED_MIME_TYPES:
                logger.error(f"Unsupported MIME type: {mime_type}")
                raise ValueError(f"Unsupported file type. Only PDF and DOCX files are supported.")
            
            file_extension = '.' + self.SUPPORTED_MIME_TYPES[mime_type]
            return self.extract_text(content, file_extension)
                
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing file {file.filename}: {str(e)}")
        finally:
            # Reset file pointer for potential reuse
            await file.seek(0)
            
    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        Extract text content from uploaded file (alias for extract_text_from_upload)
        
        This method provides compatibility with the refactored application structure.
        
        Args:
            file (UploadFile): Uploaded file object
        
        Returns:
            str: Extracted text content
            
        Raises:
            HTTPException: If file processing fails
        """
        return await self.extract_text_from_upload(file)
    
    def _extract_from_pdf(self, content: bytes, filename: str) -> str:
        """
        Extract text from PDF content
        
        Args:
            content (bytes): PDF file content
            filename (str): Original filename for logging
            
        Returns:
            str: Extracted text
        """
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            if len(reader.pages) == 0:
                logger.warning(f"PDF file {filename} has no pages")
                return ""
                
            logger.info(f"Extracting text from PDF: {filename} with {len(reader.pages)} pages")
            
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if not page_text:
                    logger.warning(f"No text extracted from page {i+1} in {filename}")
                text += page_text + "\n"
                
            extracted_text = text.strip()
            logger.info(f"Successfully extracted {len(extracted_text)} characters from {filename}")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {filename}: {str(e)}")
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_from_docx(self, content: bytes, filename: str) -> str:
        """
        Extract text from DOCX content
        
        Args:
            content (bytes): DOCX file content
            filename (str): Original filename for logging
            
        Returns:
            str: Extracted text
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            if len(doc.paragraphs) == 0:
                logger.warning(f"DOCX file {filename} has no paragraphs")
                return ""
                
            logger.info(f"Extracting text from DOCX: {filename} with {len(doc.paragraphs)} paragraphs")
            
            # Extract text from paragraphs
            text = ""
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            extracted_text = text.strip()
            logger.info(f"Successfully extracted {len(extracted_text)} characters from {filename}")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {filename}: {str(e)}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
